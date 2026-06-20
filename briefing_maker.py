# -*- coding: utf-8 -*-
"""
briefing_maker.py
=================
[월별 이슈브리핑 + 모니터링 결과 자동 생성기]

GitHub Actions에서 월(예: 202601)을 입력받아 실행됩니다.

전체 흐름:
  1) 구글 시트에서 해당 월 데이터 호출 (연관높음 + 단순관련)
  2) Gemini가 ① TOP N 선별 → ② 총평 작성 → ③ 법령별 상세 작성
  3) 이슈브리핑(.docx) + 모니터링 결과(.xlsx) 생성  ← 우리가 디자인한 그대로
  4) 두 파일을 웹훅(Make.com)으로 첨부 발송 → 공단 메일함

[필요 환경변수] (GitHub Secrets에 저장)
  GCP_SA_JSON      : 구글 서비스계정 인증 JSON
  GOOGLE_SHEET_ID  : monitor 구글시트 KEY
  GEMINI_API_KEY   : Gemini API 키
  LLM_MODEL        : (선택) 모델명. 기본 gemini-2.5-pro
  BRIEFING_WEBHOOK_URL : (선택) 이슈브리핑 전용 Make.com 웹훅. 비우면 메일 미발송(로컬 테스트용)
                         ※ monitor 일일 알림 웹훅(WEBHOOK_URL)과는 별개의 새 웹훅
  TARGET_MONTH     : 생성할 월 (예: 202601). 워크플로우 입력값으로 주입.
"""

import os
import io
import re
import json
import time
import requests
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager
from collections import Counter

# 구글 시트 인증은 기존 코어 모듈을 그대로 재사용
from hrdk_law_core.sheets import get_sheet_client
# Gemini 호출도 기존 모델 추상화 모듈 재사용 (모델 교체 가능)
from hrdk_law_core.llm_client import get_llm_client

# 프롬프트는 별도 파일에서
from briefing_prompts import (
    PERSONA, SELECT_TOP_PROMPT, FOREWORD_PROMPT, DETAIL_PROMPT
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from openpyxl import Workbook
from openpyxl.styles import Font as XLFont, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# ============================================================
# 설정값
# ============================================================
TARGET_MONTH = os.environ.get("TARGET_MONTH", "").strip()  # 예: 202601
GCP_SA_JSON = os.environ.get("GCP_SA_JSON")
GOOGLE_SHEET_ID = os.environ.get("GOOGLE_SHEET_ID")
WEBHOOK_URL = os.environ.get("BRIEFING_WEBHOOK_URL")  # 이슈브리핑 전용 웹훅 (monitor 일일 알림과 분리)

TOP_N = 5  # 이슈브리핑에 담을 핵심 법령 수 (기본 5건)

# 디자인 색상 (우리가 정한 공단 네이비 톤)
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x5A, 0x88)
GRAY = RGBColor(0x59, 0x59, 0x59)
DARKGRAY = RGBColor(0x40, 0x40, 0x40)
FONT = "맑은 고딕"

SHEET_HIGH = "연관 높은 법령"
SHEET_SIMPLE = "국가기술자격 관계 법령(단순 관련)"


# ============================================================
# [1] 구글 시트에서 해당 월 데이터 호출
# ============================================================
def norm_date(v):
    s = "".join(c for c in str(v) if c.isdigit())
    return s[:8] if len(s) >= 8 else ""


def fetch_month_data(target_month):
    """구글 시트에서 target_month(YYYYMM) 데이터를 가져온다.
    반환: (high, simple, total_laws)
      - high   : 활용 높은 법령 목록
      - simple : 단순 관련 법령 목록
      - total_laws : 그달 전체 시행 법령 수 (총괄현황표의 '총 검토건수' 합계)

    ※ 법제처 API를 따로 부르지 않는다.
      monitor가 매일 검토하며 총괄현황표에 쌓아둔 '총 검토건수'가 곧 전체 시행 법령 수다.
      (구글 시트가 우리의 영구 스토리지이므로, 모든 통계를 여기서 가져온다.)
    """
    print(f"📥 [1단계] 구글 시트에서 {target_month} 데이터 호출 중...")
    # get_sheet_client는 (client, spreadsheet) 튜플 반환 → 두 번째만 사용
    _, ss = get_sheet_client(GCP_SA_JSON, GOOGLE_SHEET_ID)

    def pull(tab_name):
        records = ss.worksheet(tab_name).get_all_records()
        return [r for r in records
                if norm_date(r.get("시행일자", "")).startswith(target_month)]

    high = pull(SHEET_HIGH)
    simple = pull(SHEET_SIMPLE)

    # 총괄현황표에서 그달 '총 검토건수' 합계 = 전체 시행 법령 수
    total_laws = sum_total_reviewed(ss, target_month)

    print(f"   → 전체 시행 {total_laws}건 / 활용 높은 {len(high)}건 / 단순 관련 {len(simple)}건")
    return high, simple, total_laws


def sum_total_reviewed(ss, target_month):
    """총괄현황표에서 target_month(YYYYMM)에 해당하는 '총 검토건수'를 모두 더한다.
    이 값이 곧 그달 전체 시행(검토 대상) 법령 수다. 실패 시 None."""
    try:
        records = ss.worksheet("총괄현황표").get_all_records()
    except Exception as e:
        print(f"   ⚠️ 총괄현황표 읽기 실패: {e}")
        return None

    # '총 검토건수' 칸 이름이 다를 수 있어 후보를 둔다 (헤더 표기 차이 대응)
    count_keys = ["총 검토건수", "총검토건수", "총 검토 건수"]
    total = 0
    found = False
    for r in records:
        # 날짜 칸도 표기 차이 대응 (시행일자/수집일자)
        date_val = r.get("시행일자") or r.get("수집일자") or ""
        if not norm_date(date_val).startswith(target_month):
            continue
        # 총 검토건수 칸 찾기
        for k in count_keys:
            if k in r:
                try:
                    total += int(float(r[k] or 0))
                    found = True
                except (ValueError, TypeError):
                    pass
                break
    return total if found else None


# ============================================================
# [2] Gemini 3단계 호출
# ============================================================
def _ask(prompt, temperature=0.2, retries=3):
    """Gemini 호출 + 재시도. 코어의 llm_client 사용."""
    llm = get_llm_client()
    for attempt in range(retries):
        try:
            return llm.generate(prompt, temperature=temperature)
        except Exception as e:
            print(f"   ⚠️ AI 호출 실패({attempt+1}/{retries}): {e}")
            time.sleep(10)
    return ""


def select_top_laws(big_laws, top_n=TOP_N):
    """대폭 증감 법령 중 핵심 N건 선별. 적으면 있는 만큼."""
    print(f"🧠 [2-1] 대폭 증감 {len(big_laws)}건 중 핵심 {top_n}건 선별 중...")
    if len(big_laws) <= top_n:
        return big_laws  # 후보가 적으면 전부 사용

    candidates = [
        {"id": i, "법령명": r.get("법령명", ""),
         "관련자격": r.get("법령 관련 국가기술자격 종목", ""),
         "활용도": r.get("활용도 분석 구분", "")}
        for i, r in enumerate(big_laws)
    ]
    prompt = SELECT_TOP_PROMPT.format(
        persona=PERSONA, top_n=top_n,
        candidates=json.dumps(candidates, ensure_ascii=False)
    )
    raw = _ask(prompt, temperature=0.0)
    # 응답에서 [숫자, 숫자, ...] 배열만 추출
    match = re.search(r"\[(.*?)\]", raw, re.DOTALL)
    if match:
        try:
            ids = json.loads(f"[{match.group(1)}]")
            picked = [big_laws[i] for i in ids if isinstance(i, int) and i < len(big_laws)]
            if picked:
                return picked[:top_n]
        except Exception:
            pass
    # AI 실패 시 안전장치: 그냥 앞 N건
    print("   ⚠️ 선별 응답 파싱 실패 → 앞 N건으로 대체")
    return big_laws[:top_n]


def make_foreword(selected, target_month):
    """이달의 총평 작성."""
    print("🧠 [2-2] 이달의 총평 작성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    summary = [{"법령명": r.get("법령명", ""),
                "내용": r.get("주요 제·개정내용", "")} for r in selected]
    prompt = FOREWORD_PROMPT.format(
        persona=PERSONA, year=year, month=month, count=len(selected),
        summary_data=json.dumps(summary, ensure_ascii=False)
    )
    text = _ask(prompt, temperature=0.3).strip()
    if not text:
        text = (f"{year}년 {month}월은 국가기술자격의 활용 기반이 전반적으로 "
                f"강화되는 추세가 확인되는 시기입니다.")
    return text


def make_details(selected):
    """선별된 법령별 상세(파급효과/배경/내용/효과) 작성."""
    print(f"🧠 [2-3] 핵심 {len(selected)}건 상세 분석 작성 중...")
    results = []
    for idx, r in enumerate(selected, 1):
        name = r.get("법령명", "")
        print(f"   📝 [{idx}/{len(selected)}] {name[:30]} ...", end=" ", flush=True)
        prompt = DETAIL_PROMPT.format(
            persona=PERSONA, law_name=name,
            enf_date=norm_date(r.get("시행일자", "")),
            dept=r.get("소관부처", ""),
            certs=r.get("법령 관련 국가기술자격 종목", ""),
            summary=r.get("주요 제·개정내용", ""),
            util_detail=r.get("활용도 분석 상세", ""),
        )
        raw = _ask(prompt, temperature=0.2)
        data = _parse_detail_json(raw)
        merged = dict(r)
        merged["impact_3lines"] = data.get("impact_3lines", [])
        merged["bg"] = data.get("bg", "")
        merged["main"] = data.get("main", "")
        merged["effect"] = data.get("effect", "")
        results.append(merged)
        print("✅")
    return results


def _parse_detail_json(raw):
    """AI 응답에서 JSON 객체 추출 (```json 펜스, 줄바꿈 등 정리)."""
    if not raw:
        return {}
    s = raw.replace("```json", "").replace("```JSON", "").replace("```", "").strip()
    s = s.replace("\n", " ").replace("\r", " ")
    match = re.search(r"\{.*\}", s, re.DOTALL)
    if not match:
        return {}
    try:
        return json.loads(match.group(0))
    except Exception:
        return {}


# ============================================================
# [3-A] 차트 생성 (부처별 TOP 5)
# ============================================================
def _korean_font():
    """GitHub Actions(Ubuntu)에 설치된 한글 폰트 경로 자동 탐색."""
    candidates = [
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJKkr-Regular.otf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return font_manager.FontProperties(fname=p)
    return None  # 없으면 기본 폰트 (한글 깨질 수 있음 → yml에서 폰트 설치)


def make_chart(high, out_path):
    print("📊 [3] 부처별 차트 생성 중...")
    fm = _korean_font()
    if fm:
        plt.rcParams["font.family"] = fm.get_name()
    plt.rcParams["axes.unicode_minus"] = False

    dept_counts = Counter(r.get("소관부처", "기타") for r in high)
    top5 = dept_counts.most_common(5)
    names = [x[0] for x in top5]
    values = [x[1] for x in top5]

    fig, ax = plt.subplots(figsize=(7, 3.8))
    bars = ax.bar(names, values, color="#2E5A88", width=0.6)
    ax.set_ylabel("관련 법령 건수", fontproperties=fm, fontsize=11)
    ax.set_title("부처별 자격 관련 법령 제·개정 현황",
                 fontproperties=fm, fontsize=13, fontweight="bold", pad=12)
    for bar in bars:
        h = bar.get_height()
        ax.text(bar.get_x() + bar.get_width() / 2, h, f"{int(h)}건",
                ha="center", va="bottom", fontproperties=fm, fontsize=10, fontweight="bold")
    ax.set_xticks(range(len(names)))
    ax.set_xticklabels(names, fontproperties=fm, fontsize=9, rotation=10)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", linestyle="--", alpha=0.4)
    if values:
        ax.set_ylim(0, max(values) * 1.18)
    plt.tight_layout()
    plt.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close()


# ============================================================
# [3-B] 이슈브리핑 docx 생성 — 우리가 디자인한 그대로
# ============================================================
def _set_cell_bg(cell, hex_color):
    """표 셀 배경색 (python-docx 기본 미지원 → XML 직접 주입)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _run(p, text, size=11, color=None, bold=False, italic=False):
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def build_briefing_docx(target_month, total_laws, related_count,
                        big_increase, foreword, issues, chart_path, out_path):
    print("📄 [4-1] 이슈브리핑 docx 생성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    style.font.size = Pt(11)

    # 페이지: US Letter, 여백 약 2.5cm
    sec = doc.sections[0]
    sec.page_width = Cm(21.59)
    sec.page_height = Cm(27.94)
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # --- 제목 ---
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "국가기술자격 관련 법령", size=20, color=NAVY, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, "Issue Briefing", size=26, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, f"〈{year}년 {month}월호〉", size=13, color=GRAY, bold=True)
    _add_bottom_border(p, NAVY)

    # 전체 시행 법령 수 표시용 (총괄현황표에서 못 읽으면 None)
    has_total = bool(total_laws) and total_laws > related_count
    total_str = f"{total_laws}건" if has_total else None

    # --- 개요 ---
    p = doc.add_paragraph(); _run(p, "■ 개요", size=12, color=BLUE, bold=True)
    overview = [
        ("조사기간", f"{year}년 {month}월 1일 ~ {month}월 말일"),
    ]
    if has_total:
        overview.append(("조사대상", f"{month}월 시행 법령 총 {total_str} (국가기술자격 관련 {related_count}건)"))
    else:
        overview.append(("조사대상", f"{month}월 시행 국가기술자격 관련 법령 {related_count}건"))
    overview.append(("주요내용", f"자격 활용도 변동 {len(issues)}개 핵심 사례"))
    for label, val in overview:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, f"✓ {label} : ", size=10.5, color=NAVY, bold=True)
        _run(p, val, size=10.5)

    # --- 모니터링 요약 ---
    p = doc.add_paragraph(); _run(p, "■ 모니터링 요약", size=12, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    if has_total:
        # 전체 수를 아는 경우: "전체 N건 중, 관련 M건"
        _run(p, f"· {month}월 전체 시행 법령 ", size=10.5)
        _run(p, total_str, size=10.5, color=NAVY, bold=True)
        _run(p, " 중, 국가기술자격 관련 법령은 ", size=10.5)
        _run(p, f"{related_count}건", size=10.5, color=NAVY, bold=True)
        _run(p, "으로 조사", size=10.5)
    else:
        # 전체 수를 모르는 경우: 관련 법령 수만
        _run(p, f"· {month}월 국가기술자격 관련 법령은 ", size=10.5)
        _run(p, f"{related_count}건", size=10.5, color=NAVY, bold=True)
        _run(p, "으로 조사", size=10.5)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "· 그중 자격 활용도가 ", size=10.5)
    _run(p, "대폭 증가", size=10.5, color=RGBColor(0xC5, 0x5A, 0x11), bold=True)
    _run(p, f"한 법령은 {big_increase}건으로 조사", size=10.5)

    # --- 총평 (회색 박스) ---
    h = doc.add_paragraph(); _run(h, "이달의 주요 정책 트렌드 및 총평", size=13, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4); p.paragraph_format.right_indent = Cm(0.4)
    _run(p, foreword, size=10.5, color=DARKGRAY)
    _shade_paragraph(p, "F2F2F2")
    _add_left_border(p, BLUE)

    # --- 차트 ---
    h = doc.add_paragraph(); _run(h, "데이터 시각화 분석", size=13, color=NAVY, bold=True)
    if os.path.exists(chart_path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(chart_path, width=Cm(12.5))

    # --- 페이지 나눔 → 요약표 ---
    doc.add_page_break()
    h = doc.add_paragraph(); _run(h, "〈 주요 제·개정 법령 요약 〉", size=12, color=NAVY, bold=True)
    _build_summary_table(doc, issues)

    # --- 페이지 나눔 → 법령별 상세 ---
    doc.add_page_break()
    h = doc.add_paragraph(); _run(h, "〈 핵심 법령 상세 분석 〉", size=12, color=NAVY, bold=True)
    for i, it in enumerate(issues, 1):
        _build_detail_card(doc, i, it)

    # 머리말/푸터
    _add_header_footer(doc)
    doc.save(out_path)


def _build_summary_table(doc, issues):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    widths = [Cm(1.2), Cm(5.5), Cm(2.0), Cm(4.0), Cm(3.5)]
    hdr = table.rows[0].cells
    for c, (txt, w) in enumerate(zip(["연번", "법령명", "시행일", "관련 자격", "핵심 내용"], widths)):
        hdr[c].width = w
        _set_cell_bg(hdr[c], "1F3864")
        para = hdr[c].paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(para, txt, size=9, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    for i, it in enumerate(issues, 1):
        cells = table.add_row().cells
        date = norm_date(it.get("시행일자", ""))
        date_fmt = f"{date[:4]}. {date[4:6]}. {date[6:8]}." if len(date) == 8 else date
        certs = _summarize_certs(it.get("법령 관련 국가기술자격 종목", ""))
        summ = str(it.get("주요 제·개정내용", ""))[:40] + "…"
        vals = [str(i), it.get("법령명", ""), date_fmt, certs, summ]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT]
        for c, (v, w, al) in enumerate(zip(vals, widths, aligns)):
            cells[c].width = w
            para = cells[c].paragraphs[0]; para.alignment = al
            _run(para, v, size=8.5)


def _build_detail_card(doc, i, it):
    # 제목 바 (네이비)
    p = doc.add_paragraph()
    _run(p, f"{i}. {it.get('법령명','')}", size=11, color=RGBColor(0xFF, 0xFF, 0xFF), bold=True)
    _shade_paragraph(p, "1F3864")
    # 메타 (연파랑)
    date = norm_date(it.get("시행일자", ""))
    date_fmt = f"{date[:4]}. {date[4:6]}. {date[6:8]}." if len(date) == 8 else date
    certs = _summarize_certs(it.get("법령 관련 국가기술자격 종목", ""))
    p = doc.add_paragraph()
    _run(p, "시행일 ", size=8.5, color=NAVY, bold=True); _run(p, f"{date_fmt}    ", size=8.5)
    _run(p, "관련 자격 ", size=8.5, color=NAVY, bold=True); _run(p, certs, size=8.5)
    _shade_paragraph(p, "D6E2F0")
    # 파급효과
    p = doc.add_paragraph(); _run(p, "▣ 자격증 파급효과", size=10, color=BLUE, bold=True)
    for k, line in enumerate(it.get("impact_3lines", []), 1):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.6)
        _run(p, f"{k}. ", size=9.5, color=NAVY, bold=True); _run(p, line, size=9.5)
    # 상세 분석
    p = doc.add_paragraph(); _run(p, "▣ 상세 분석", size=10, color=BLUE, bold=True)
    for lbl, val in [("□ 추진배경", it.get("bg", "")),
                     ("□ 주요 개정내용", it.get("main", "")),
                     ("□ 자격증 기대효과", it.get("effect", ""))]:
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
        _run(p, lbl, size=9.5, color=DARKGRAY, bold=True)
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.8)
        _run(p, val, size=9.5, color=DARKGRAY)
    # 개정 조문 (회색, 본문과 분리)
    p = doc.add_paragraph(); _run(p, "▣ 개정 조문", size=10, color=BLUE, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    jomun = str(it.get("근거조문", "") or "(해당 없음)")
    _run(p, jomun, size=8.5, color=GRAY, italic=True)
    _shade_paragraph(p, "F2F2F2")
    # 구분선
    p = doc.add_paragraph(); _add_bottom_border(p, RGBColor(0xBF, 0xBF, 0xBF), dashed=True)


def _summarize_certs(cert_string):
    certs = [c.strip() for c in str(cert_string).split(",") if c.strip()]
    if len(certs) > 1:
        return f"{certs[0]} 등 {len(certs)}개 종목"
    return certs[0] if certs else "해당 없음"


# --- 단락 음영/테두리 유틸 (XML 직접 조작) ---
def _shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _add_bottom_border(p, color, dashed=False):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "dashed" if dashed else "single")
    bottom.set(qn("w:sz"), "6"); bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(bottom); pPr.append(pbdr)


def _add_left_border(p, color):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "24"); left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pbdr.append(left); pPr.append(pbdr)


def _add_header_footer(doc):
    sec = doc.sections[0]
    hdr = sec.header.paragraphs[0]; hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(hdr, "국가기술자격 관련 법령 Issue Briefing", size=7.5, color=GRAY)
    ftr = sec.footer.paragraphs[0]; ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(ftr, "한국산업인력공단", size=7.5, color=GRAY)


# ============================================================
# [3-C] 모니터링 결과 xlsx 생성 — 근거조문 + 하이퍼링크
# ============================================================
def build_monitor_xlsx(target_month, total_laws, high, simple, out_path):
    print("📊 [4-2] 모니터링 결과 xlsx 생성 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    YM = f"{year}년 {month}월"

    XLNAVY = "1F3864"; XLBLUE = "2E5A88"; XLLIGHTBLUE = "D6E2F0"
    YELLOW = "FFF2CC"; LINKBLUE = "0563C1"; HGRAY = "404040"
    thin = Side(style="thin", color="BFBFBF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left = Alignment(horizontal="left", vertical="center", wrap_text=True)

    def hdr(cell, fill=XLNAVY, color="FFFFFF"):
        cell.font = XLFont(name=FONT, bold=True, color=color, size=10)
        cell.fill = PatternFill("solid", fgColor=fill); cell.alignment = center; cell.border = border

    def body(cell, fill=None, align=left, size=9):
        cell.font = XLFont(name=FONT, size=size); cell.alignment = align; cell.border = border
        if fill:
            cell.fill = PatternFill("solid", fgColor=fill)

    def law_url(nm):
        return f"https://www.law.go.kr/법령/{nm}"

    wb = Workbook()

    # 시트1: 요약
    ws = wb.active; ws.title = "모니터링 요약"
    ws.merge_cells("A1:C1")
    ws["A1"] = f"{YM} 국가기술자격 관련 법령 제·개정사항 모니터링 결과"
    ws["A1"].font = XLFont(name=FONT, bold=True, size=14, color=XLNAVY)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30
    related_total = len(high) + len(simple)
    has_total = bool(total_laws) and total_laws > related_total
    ws.merge_cells("A2:C2")
    total_label = f"총 {total_laws}건" if has_total else f"국가기술자격 관련 {related_total}건"
    ws["A2"] = f"※ {YM} 시행 법령 : {total_label}"
    ws["A2"].font = XLFont(name=FONT, size=10, color=HGRAY)

    util = Counter(r.get("활용도 분석 구분", "") for r in high)
    stats = [
        ("구분", "건수", "비고", True),
        ("총 시행 법령", total_laws if has_total else "-", "당월 시행된 전체 법령", False),
        ("국가기술자격 관계 법령", related_total, "자격 관련 있는 법령 (아래 합계)", False),
        ("  ① 활용·관련 높은 법령", len(high), "자격 활용도에 영향", False),
        ("      · 대폭 증가", util.get("대폭 증가", 0), "자격 수요 크게 증가", False),
        ("      · 소폭 증가", util.get("소폭 증가", 0), "자격 수요 소폭 증가", False),
        ("      · 현상 유지", util.get("현상 유지", 0), "변동 미미", False),
        ("  ② 단순 관련 법령", len(simple), "자격 언급되나 활용도 변동 적음", False),
    ]
    for i, (label, cnt, note, is_h) in enumerate(stats):
        r = 4 + i
        ws[f"A{r}"], ws[f"B{r}"], ws[f"C{r}"] = label, cnt, note
        if is_h:
            for col in "ABC":
                hdr(ws[f"{col}{r}"])
        else:
            fill = XLLIGHTBLUE if label == "국가기술자격 관계 법령" else None
            body(ws[f"A{r}"], fill=fill, align=left, size=10)
            body(ws[f"B{r}"], fill=fill or (YELLOW if "①" in label else None), align=center, size=10)
            body(ws[f"C{r}"], fill=fill, align=left, size=9)
            ws[f"A{r}"].font = XLFont(name=FONT, size=10,
                                      bold=("관계 법령" in label or "①" in label or "②" in label))
            ws[f"B{r}"].font = XLFont(name=FONT, size=10, bold=True, color=XLNAVY)
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 38

    # 시트2: 활용 높은 법령
    ws2 = wb.create_sheet("활용 높은 법령")
    ws2.merge_cells("A1:H1")
    ws2["A1"] = f"1. 국가기술자격 활용·관련 높은 법령 : {len(high)}건"
    ws2["A1"].font = XLFont(name=FONT, bold=True, size=12, color=XLNAVY)
    ws2.row_dimensions[1].height = 24
    h_head = ["연번", "법령명", "시행일자", "소관부처", "주요 제·개정내용",
              "관련 국가기술자격 종목", "활용도 구분", "근거 조문"]
    h_w = [6, 30, 12, 15, 48, 38, 11, 40]
    for c, (t, w) in enumerate(zip(h_head, h_w), 1):
        hdr(ws2.cell(row=2, column=c, value=t))
        ws2.column_dimensions[get_column_letter(c)].width = w
    order = {"대폭 증가": 0, "대폭 감소": 1, "소폭 증가": 2, "소폭 감소": 3, "현상 유지": 4}
    for i, r in enumerate(sorted(high, key=lambda x: order.get(x.get("활용도 분석 구분", ""), 9)), 1):
        row = 2 + i
        d = norm_date(r.get("시행일자", ""))
        df = f"{d[:4]}-{d[4:6]}-{d[6:8]}" if len(d) == 8 else d
        is_big = r.get("활용도 분석 구분", "") in ("대폭 증가", "대폭 감소")
        fill = YELLOW if is_big else None
        body(ws2.cell(row=row, column=1, value=i), fill=fill, align=center)
        c2 = ws2.cell(row=row, column=2, value=r.get("법령명", ""))
        c2.hyperlink = law_url(r.get("법령명", ""))
        c2.font = XLFont(name=FONT, size=9, color=LINKBLUE, underline="single")
        c2.alignment = left; c2.border = border
        if fill:
            c2.fill = PatternFill("solid", fgColor=fill)
        body(ws2.cell(row=row, column=3, value=df), fill=fill, align=center)
        body(ws2.cell(row=row, column=4, value=r.get("소관부처", "")), fill=fill, align=left)
        body(ws2.cell(row=row, column=5, value=str(r.get("주요 제·개정내용", ""))[:200]), fill=fill, align=left)
        body(ws2.cell(row=row, column=6, value=str(r.get("법령 관련 국가기술자격 종목", ""))[:150]), fill=fill, align=left)
        body(ws2.cell(row=row, column=7, value=r.get("활용도 분석 구분", "")), fill=fill, align=center)
        body(ws2.cell(row=row, column=8, value=str(r.get("근거조문", ""))[:200]), fill=fill, align=left)
        ws2.row_dimensions[row].height = 48
    ws2.freeze_panes = "A3"

    # 시트3: 단순 관련 법령
    ws3 = wb.create_sheet("단순 관련 법령")
    ws3.merge_cells("A1:G1")
    ws3["A1"] = f"2. 국가기술자격 관계 법령(단순 관련) 제·개정 사항 : {len(simple)}건"
    ws3["A1"].font = XLFont(name=FONT, bold=True, size=12, color=XLNAVY)
    ws3.row_dimensions[1].height = 24
    s_head = ["연번", "법령명", "시행일자", "소관부처", "주요 제·개정이유",
              "관련 국가기술자격 종목", "근거 조문"]
    s_w = [6, 30, 12, 15, 52, 42, 38]
    for c, (t, w) in enumerate(zip(s_head, s_w), 1):
        hdr(ws3.cell(row=2, column=c, value=t), fill=XLBLUE)
        ws3.column_dimensions[get_column_letter(c)].width = w
    for i, r in enumerate(simple, 1):
        row = 2 + i
        d = norm_date(r.get("시행일자", ""))
        df = f"{d[:4]}-{d[4:6]}-{d[6:8]}" if len(d) == 8 else d
        body(ws3.cell(row=row, column=1, value=i), align=center)
        c2 = ws3.cell(row=row, column=2, value=r.get("법령명", ""))
        c2.hyperlink = law_url(r.get("법령명", ""))
        c2.font = XLFont(name=FONT, size=9, color=LINKBLUE, underline="single")
        c2.alignment = left; c2.border = border
        body(ws3.cell(row=row, column=3, value=df), align=center)
        body(ws3.cell(row=row, column=4, value=r.get("소관부처", "")), align=left)
        body(ws3.cell(row=row, column=5, value=str(r.get("주요 제·개정내용", ""))[:250]), align=left)
        body(ws3.cell(row=row, column=6, value=str(r.get("법령 관련 국가기술자격 종목", ""))[:150]), align=left)
        body(ws3.cell(row=row, column=7, value=str(r.get("근거조문", ""))[:200]), align=left)
        ws3.row_dimensions[row].height = 48
    ws3.freeze_panes = "A3"

    wb.save(out_path)


# ============================================================
# [4] 웹훅으로 두 파일 첨부 발송 (Make.com → 메일)
# ============================================================
def send_via_webhook(target_month, docx_path, xlsx_path, stats):
    if not WEBHOOK_URL:
        print("ℹ️ BRIEFING_WEBHOOK_URL이 없어 메일 발송을 건너뜁니다. (로컬 테스트 모드)")
        return
    print("📧 [5] 웹훅으로 보고서 발송 중...")
    year, month = target_month[:4], str(int(target_month[4:6]))
    data = {
        "system": "law-monitor-briefing",
        "source": "briefing",
        "subject": f"[이슈브리핑] {year}년 {month}월호 국가기술자격 관련 법령",
        "month": f"{year}년 {month}월",
        "total": str(stats.get("total", "")),
        "related": str(stats.get("related", "")),
        "big": str(stats.get("big", "")),
    }
    files = {}
    fh1 = open(docx_path, "rb"); fh2 = open(xlsx_path, "rb")
    try:
        files = {
            "file1": (os.path.basename(docx_path), fh1,
                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "file2": (os.path.basename(xlsx_path), fh2,
                      "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        }
        requests.post(WEBHOOK_URL, data=data, files=files, timeout=60)
        print("   ✅ 발송 완료")
    except Exception as e:
        print(f"   ⚠️ 발송 실패: {e}")
    finally:
        fh1.close(); fh2.close()


# ============================================================
# 메인
# ============================================================
def main():
    if not TARGET_MONTH or len(TARGET_MONTH) != 6:
        raise SystemExit("❌ TARGET_MONTH(YYYYMM)이 필요합니다. 예: 202601")

    print(f"🚀 {TARGET_MONTH} 이슈브리핑 생성 시작\n" + "=" * 50)

    # 1) 데이터 (전체 시행 법령 수도 총괄현황표에서 함께 가져옴)
    high, simple, total_laws = fetch_month_data(TARGET_MONTH)
    if not high and not simple:
        raise SystemExit(f"❌ {TARGET_MONTH} 데이터가 시트에 없습니다.")

    related_count = len(high) + len(simple)
    util = Counter(r.get("활용도 분석 구분", "") for r in high)
    big_increase = util.get("대폭 증가", 0) + util.get("대폭 감소", 0)

    # 2) AI: 선별 → 총평 → 상세
    big_laws = [r for r in high if r.get("활용도 분석 구분", "") in ("대폭 증가", "대폭 감소")]
    selected = select_top_laws(big_laws, TOP_N)
    foreword = make_foreword(selected, TARGET_MONTH)
    issues = make_details(selected)

    # 3) 차트 + 두 산출물
    chart_path = "/tmp/chart.png"
    make_chart(high, chart_path)

    docx_path = f"/tmp/이슈브리핑_{TARGET_MONTH}.docx"
    xlsx_path = f"/tmp/모니터링결과_{TARGET_MONTH}.xlsx"
    build_briefing_docx(TARGET_MONTH, total_laws, related_count, big_increase,
                        foreword, issues, chart_path, docx_path)
    build_monitor_xlsx(TARGET_MONTH, total_laws, high, simple, xlsx_path)

    # GitHub Actions가 가져갈 수 있게 현재 폴더에도 복사
    import shutil
    out_dir = os.environ.get("OUTPUT_DIR", ".")
    os.makedirs(out_dir, exist_ok=True)
    final_docx = os.path.join(out_dir, os.path.basename(docx_path))
    final_xlsx = os.path.join(out_dir, os.path.basename(xlsx_path))
    shutil.copy(docx_path, final_docx)
    shutil.copy(xlsx_path, final_xlsx)
    print(f"\n📁 생성 완료: {final_docx}, {final_xlsx}")

    # 4) 발송
    send_via_webhook(TARGET_MONTH, final_docx, final_xlsx,
                     {"total": total_laws, "related": related_count, "big": big_increase})

    print("=" * 50 + "\n✨ 이슈브리핑 생성 완료!")


if __name__ == "__main__":
    main()
